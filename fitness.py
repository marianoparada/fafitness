import streamlit as st
import random
import time
from docx import Document
from docx.shared import Inches
import io
import pandas as pd

st.set_page_config(layout="wide")  # Configura la página para usar todo el ancho
hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            #header {visibility: hidden;}
            </style>
            """
#st.markdown(hide_st_style, unsafe_allow_html=True)
# Estilos personalizados para los botones
button_style = """
    <style>
    div.stButton > button {
        width: 100%;
        height: 100px;
        font-size: 24px;
        padding: 20px;
    }
    </style>
    """
st.markdown(button_style, unsafe_allow_html=True)

def clear_page():
    """Limpia el contenido de la página actual"""
    st.empty()
    
def load_data(pagina):
    try:
        sheet_id = '1bSlKMIBnXdI0R8JrYxzp622HuwYp2Q7xc2Fz3apnHlw'
        df = pd.read_csv(f'https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={pagina}')
        # Guardar en caché local
        df.to_csv('local_cache.csv', index=False)
        return df
    except Exception as e:
        st.warning("No se pudo cargar datos de Google Sheets. Intentando cargar desde caché local.")
        try:
            return pd.read_csv('local_cache.csv')
        except FileNotFoundError:
            st.error("No se pudo cargar datos. Por favor, verifica tu conexión a internet.")
            return pd.DataFrame()



#df=load_data('entrenahoy')


def mostrar_link():
    links=load_data('links')
    random_row = links.sample(n=1).iloc[0]
    # Mostrar el comentario y propietario
    
    # Mostrar el video en Streamlit
    video_link = random_row['link']
    st.video(video_link)
    st.title(f"{random_row['comentario']} - ({random_row['propietario']})")
    # Mostrar el link para abrir en otra ventana
    st.markdown(f"[Abrir video en otra ventana]({video_link})")

def mostrar_banner():
    banners = ["entrena.png"]  # Agrega todos tus banners aquí
    # Selecciona un banner al azar
    banner_aleatorio = random.choice(banners)
    # Muestra el banner al azar
    st.image(banner_aleatorio, use_column_width=True)

def mostrar_logo():
    if 'logo_mostrado' not in st.session_state:
        st.session_state.logo_mostrado = False

    if not st.session_state.logo_mostrado:
        left_co, cent_co, last_co = st.columns(3)
        with cent_co:
            st.image('logo.jpg')
            st.session_state.logo_mostrado = True

def guardar_rutina_word(rutina):
    doc = Document()
    doc.add_heading('Fitness - Rutina Personalizada', 0)

    vuelta_actual = 1
    for grupo, ejercicio, duracion, instruccion, vuelta, num_ejercicio, total_ejercicios in rutina:
        if vuelta != vuelta_actual:
            doc.add_paragraph('---')
            doc.add_paragraph(f'Vuelta {vuelta}')
            vuelta_actual = vuelta
        if grupo == "Descanso":
            if ejercicio == "Descanso entre vueltas":
                doc.add_paragraph(f'Descanso entre vueltas: {duracion} segundos')
            else:
                doc.add_paragraph(f'Descanso: {duracion} segundos')
        else:
            doc.add_paragraph(f'Ejercicio {num_ejercicio}/{total_ejercicios}: {grupo} - {ejercicio}: {duracion} segundos')
            doc.add_paragraph(f'Instrucción: {instruccion}')

    # Guardar el documento en un objeto BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Botón de descarga
    st.download_button(
        label="Guardar Rutina",
        data=buffer,
        file_name="FITNESS_Rutina.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="descargar_rutina"
    )
# Configuración de la página
#st.set_page_config(page_title="Fitness - Generador de Rutinas", layout="wide")


ejercicios = {
    "💪 Brazos": [
        ("Curl de bíceps (mancuernas)", "De pie, sostenga una mancuerna en cada mano con los brazos extendidos. Doble los codos para levantar las mancuernas hacia los hombros, luego baje lentamente."),
        ("Extensiones de tríceps (mancuernas)", "Siéntese o póngase de pie con una mancuerna sostenida con ambas manos sobre la cabeza. Baje la mancuerna detrás de la cabeza doblando los codos, luego extienda los brazos."),
        ("Flexiones de brazos", "Acuéstese boca abajo con las manos apoyadas en el suelo al ancho de los hombros. Empuje el cuerpo hacia arriba manteniendo el cuerpo recto, luego baje."),
        ("Fondos de tríceps (silla)", "Siéntese en el borde de una silla, coloque las manos a los lados y deslice el cuerpo hacia adelante. Baje el cuerpo doblando los codos y luego empuje hacia arriba."),
        ("Curl martillo (mancuernas)", "De pie, sostenga una mancuerna en cada mano con las palmas hacia el cuerpo. Doble los codos para levantar las mancuernas hacia los hombros, luego baje lentamente."),
        ("Flexiones diamante", "Acuéstese boca abajo con las manos juntas formando un diamante debajo del pecho. Empuje el cuerpo hacia arriba manteniendo el cuerpo recto, luego baje."),
        ("Curl de bíceps concentrado (mancuernas)", "Siéntese en un banco, sostenga una mancuerna en una mano y apoye el codo en el muslo. Doble el codo para levantar la mancuerna, luego baje lentamente."),
        ("Extensiones de tríceps con cuerda (máquina)", "De pie, sostenga la cuerda conectada a la máquina de polea alta con ambas manos. Baje la cuerda doblando los codos, luego extienda los brazos."),
        ("Curl de bíceps en banco inclinado (mancuernas)", "Siéntese en un banco inclinado, sostenga una mancuerna en cada mano con los brazos colgando hacia abajo. Doble los codos para levantar las mancuernas hacia los hombros, luego baje lentamente."),
        ("Extensiones de tríceps sobre la cabeza (mancuernas)", "Siéntese o póngase de pie con una mancuerna sostenida con ambas manos sobre la cabeza. Baje la mancuerna detrás de la cabeza doblando los codos, luego extienda los brazos.")
    ],
    "🦵 Piernas": [
        ("Sentadillas", "De pie con los pies separados al ancho de los hombros, baje el cuerpo como si fuera a sentarse en una silla invisible, manteniendo el pecho erguido. Luego, vuelva a la posición inicial."),
        ("Estocadas", "De pie, dé un paso largo hacia adelante con una pierna. Baje el cuerpo hasta que ambas rodillas estén dobladas en ángulos de 90 grados. Empuje hacia atrás para volver a la posición inicial y alterne las piernas."),
        ("Peso muerto (barra)", "De pie con los pies separados al ancho de los hombros, sostenga una barra frente a los muslos. Inclínese hacia adelante desde las caderas, manteniendo la espalda recta, hasta que la barra llegue a las espinillas. Luego, vuelva a la posición inicial."),
        ("Elevaciones de talones", "De pie, levántese sobre los dedos de los pies, manteniendo las rodillas rectas. Baje lentamente los talones hacia el suelo."),
        ("Puente de glúteos", "Acuéstese boca arriba con las rodillas dobladas y los pies apoyados en el suelo. Levante las caderas hacia el techo, apretando los glúteos, luego baje lentamente."),
        ("Prensa de pierna (máquina)", "Siéntese en la máquina de prensa de pierna con los pies apoyados en la plataforma. Empuje la plataforma hacia adelante hasta que las piernas estén extendidas, luego baje lentamente."),
        ("Extensiones de pierna (máquina)", "Siéntese en la máquina de extensión de pierna con las piernas debajo de la almohadilla. Extienda las piernas hacia adelante, luego baje lentamente."),
        ("Curl de pierna (máquina)", "Acuéstese boca abajo en la máquina de curl de pierna con los tobillos debajo de la almohadilla. Doble las rodillas para levantar la almohadilla hacia los glúteos, luego baje lentamente."),
        ("Sentadillas búlgaras (mancuernas)", "De pie, coloque un pie en un banco detrás de usted y sostenga una mancuerna en cada mano. Baje el cuerpo doblando la rodilla de la pierna delantera, luego empuje hacia arriba."),
        ("Saltos al cajón", "De pie frente a un cajón, salte con ambos pies para aterrizar en el cajón, luego baje de un paso.")
    ],
    "🏋️ Abdominales": [
        ("Abdominales", "Acuéstese boca arriba con las rodillas dobladas y los pies apoyados en el suelo. Levante la parte superior del cuerpo hacia las rodillas, luego baje lentamente."),
        ("Plancha", "Acuéstese boca abajo con los antebrazos apoyados en el suelo y los codos debajo de los hombros. Levante el cuerpo manteniéndolo recto desde la cabeza hasta los pies."),
        ("Giros rusos", "Siéntese en el suelo con las rodillas dobladas y los pies elevados. Sostenga un peso con ambas manos y gire el torso de un lado a otro."),
        ("Elevaciones de piernas", "Acuéstese boca arriba con las piernas rectas. Levante las piernas hacia el techo hasta que los glúteos se despeguen del suelo, luego baje lentamente."),
        ("Bicicleta", "Acuéstese boca arriba con las manos detrás de la cabeza y las piernas levantadas. Alterne llevando el codo hacia la rodilla opuesta mientras extiende la otra pierna."),
        ("Plancha lateral", "Acuéstese de lado con un antebrazo apoyado en el suelo y el codo debajo del hombro. Levante el cuerpo manteniéndolo recto desde la cabeza hasta los pies."),
        ("Crunch inverso", "Acuéstese boca arriba con las piernas dobladas y los pies elevados. Levante las caderas hacia el techo, luego baje lentamente."),
        ("Escaladores", "Empiece en posición de flexión de brazos. Lleve una rodilla hacia el pecho, luego alterne rápidamente las piernas."),
        ("V-ups", "Acuéstese boca arriba con los brazos extendidos sobre la cabeza. Levante simultáneamente las piernas y el torso para tocarse los pies, luego baje lentamente."),
        ("Plancha con elevación de brazo", "Empiece en posición de plancha. Levante un brazo extendido hacia adelante, luego alterne los brazos.")
    ],
    "🏋️ Hombros": [
        ("Press militar (mancuernas)", "De pie o sentado, sostenga una mancuerna en cada mano a la altura de los hombros. Empuje las mancuernas hacia arriba hasta que los brazos estén extendidos, luego baje lentamente."),
        ("Elevaciones laterales (mancuernas)", "De pie, sostenga una mancuerna en cada mano a los lados. Levante los brazos hacia los lados hasta que estén a la altura de los hombros, luego baje lentamente."),
        ("Face pulls (cable)", "De pie, sostenga la cuerda conectada a la máquina de polea alta con ambas manos. Tire de la cuerda hacia la cara, manteniendo los codos altos."),
        ("Press Arnold (mancuernas)", "Sentado, sostenga una mancuerna en cada mano a la altura de los hombros con las palmas hacia adentro. Gire las palmas hacia afuera mientras empuja las mancuernas hacia arriba, luego baje lentamente."),
        ("Elevaciones frontales (mancuernas)", "De pie, sostenga una mancuerna en cada mano frente a los muslos. Levante los brazos hacia adelante hasta que estén a la altura de los hombros, luego baje lentamente."),
        ("Remo al mentón (barra)", "De pie, sostenga una barra con las manos juntas frente a los muslos. Levante la barra hacia el mentón, manteniendo los codos altos, luego baje lentamente."),
        ("Encogimientos de hombros (mancuernas)", "De pie, sostenga una mancuerna en cada mano a los lados. Levante los hombros hacia las orejas, luego baje lentamente."),
        ("Elevaciones posteriores (mancuernas)", "De pie, inclínese hacia adelante desde las caderas con una mancuerna en cada mano. Levante los brazos hacia los lados hasta que estén a la altura de los hombros, luego baje lentamente."),
        ("Press de hombros con barra", "Sentado o de pie, sostenga una barra a la altura de los hombros con las palmas hacia adelante. Empuje la barra hacia arriba hasta que los brazos estén extendidos, luego baje lentamente."),
        ("Remo con barra en T", "De pie, inclínese hacia adelante desde las caderas y sostenga una barra en T con ambas manos. Tire de la barra hacia el pecho, luego baje lentamente.")
    ],
    "🏃 Aeróbico": [
        ("Saltar la soga", "Salte con ambos pies mientras gira la soga por encima y por debajo del cuerpo."),
        ("Burpees", "Desde una posición de pie, agáchese y coloque las manos en el suelo. Salte los pies hacia atrás para llegar a una posición de flexión, haga una flexión, salte los pies hacia adelante y levántese saltando."),
        ("Escaladores", "Empiece en posición de flexión de brazos. Lleve una rodilla hacia el pecho, luego alterne rápidamente las piernas."),
        ("Correr", "Corra a un ritmo constante durante un período de tiempo o distancia."),
        ("Piques", "Corra a máxima velocidad durante una distancia corta, luego descanse y repita."),
        ("Saltos de tijera", "Desde una posición de pie, salte abriendo las piernas y levantando los brazos por encima de la cabeza, luego vuelva a la posición inicial."),
        ("Rodillas altas", "Corra en el lugar llevando las rodillas lo más alto posible."),
        ("Trote en el lugar", "Corra suavemente en el lugar, levantando los pies del suelo."),
        ("Escaladores", "Empiece en posición de flexión de brazos. Lleve una rodilla hacia el pecho, luego alterne rápidamente las piernas."),
        ("Saltos al banco", "De pie frente a un banco, salte con ambos pies para aterrizar en el banco, luego baje de un paso.")
    ]
}

# Distribución de ejercicios según la priorización (sin cambios)
distribucion_ejercicios = {
    "Tren superior": {"💪 Brazos": 3, "🏋️ Hombros": 2, "🏋️ Abdominales": 1, "🦵 Piernas": 1, "🏃 Aeróbico": 1},
    "Zona media": {"🏋️ Abdominales": 3, "💪 Brazos": 2, "🦵 Piernas": 1, "🏋️ Hombros": 1, "🏃 Aeróbico": 1},
    "Tren inferior": {"🦵 Piernas": 3, "🏋️ Abdominales": 2, "💪 Brazos": 1, "🏋️ Hombros": 1, "🏃 Aeróbico": 1},
    "Aeróbico": {"🏃 Aeróbico": 3, "🦵 Piernas": 2, "💪 Brazos": 1, "🏋️ Abdominales": 1, "🏋️ Hombros": 1},
    "Aleatoria": {"🏃 Aeróbico": 3, "🦵 Piernas": 2, "💪 Brazos": 1, "🏋️ Abdominales": 1, "🏋️ Hombros": 1}
}

import pandas as pd
import random

def generar_rutina(duracion, tiempo_descanso, vueltas):
    df=load_data('entrenahoy')
    rutina_base = []
    categorias = df['Categoría'].unique()
    
    for categoria in categorias:
        ejercicios_categoria = df[df['Categoría'] == categoria]
        ejercicios_seleccionados = ejercicios_categoria.sample(n=2)
        
        for _, ejercicio in ejercicios_seleccionados.iterrows():
            if duracion == "Ambos aleatorios":
                tiempo = random.choice([30, 40])
            else:
                tiempo = int(duracion.split()[0])
            
            rutina_base.append((
                ejercicio['Categoría'],
                f"{ejercicio['Icono Unicode']} {ejercicio['Ejercicio']}",
                tiempo,
                ejercicio['Descripción']
            ))
    
    rutina_completa = []
    for vuelta in range(1, vueltas + 1):
        for i, ejercicio in enumerate(rutina_base, 1):
            categoria, nombre_ejercicio, tiempo, instruccion = ejercicio
            rutina_completa.append((categoria, nombre_ejercicio, tiempo, instruccion, vuelta, i, len(rutina_base)))
            if i < len(rutina_base):
                rutina_completa.append(("Descanso", "Descanso entre ejercicios", tiempo_descanso, "Toma un breve descanso antes del siguiente ejercicio", vuelta, None, None))
        if vuelta < vueltas:
            rutina_completa.append(("Descanso", "Descanso entre vueltas", tiempo_descanso * 5, "Toma un descanso más largo antes de la siguiente vuelta", vuelta, None, None))
    
    return rutina_completa

def mostrar_rutina(rutina):
    st.title("🏋️ Fitness - Rutina Personalizada")
    vuelta_actual = 1
    for grupo, ejercicio, duracion, instruccion, vuelta, num_ejercicio, total_ejercicios in rutina:
        if vuelta != vuelta_actual:
            st.write("---")
            st.write(f"Vuelta {vuelta}")
            vuelta_actual = vuelta
        if grupo == "Descanso":
            if ejercicio == "Descanso entre vueltas":
                st.write(f"Descanso entre vueltas: {duracion} segundos")
            else:
                st.write(f"Descanso: {duracion} segundos")
        else:
            st.write(f"Ejercicio {num_ejercicio}/{total_ejercicios}: {grupo} - {ejercicio}: {duracion} segundos")
            st.write(f"Instrucción: {instruccion}")

def temporizador(rutina):
    #st.title("🏋️ FITNESS - A entrenar !")
    st.image("entrena2.png", use_column_width=True)
    if 'ejercicio_actual' not in st.session_state:
        st.session_state.ejercicio_actual = 0
        st.session_state.tiempo_restante = rutina[0][2]
    
    grupo, ejercicio, duracion, instruccion, vuelta, num_ejercicio, total_ejercicios = rutina[st.session_state.ejercicio_actual]
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        if grupo != "Descanso":
            st.markdown(f"<h3 style='text-align: center;'>Vuelta {vuelta}/{st.session_state.vueltas} - Ejercicio {num_ejercicio}/{total_ejercicios}</h3>", unsafe_allow_html=True)
        
        if grupo == "Descanso":
            if ejercicio == "Descanso entre vueltas":
                st.markdown("<h1 style='text-align: center;'>Descanso entre vueltas</h1>", unsafe_allow_html=True)
            else:
                st.markdown("<h1 style='text-align: center;'>Descanso</h1>", unsafe_allow_html=True)
            # Mostrar el próximo ejercicio
            proximo_indice = st.session_state.ejercicio_actual + 1
            if proximo_indice < len(rutina):
                proximo_ejercicio = rutina[proximo_indice][1]
                st.markdown(f"<h3 style='text-align: center;'>Próximo ejercicio: {proximo_ejercicio}</h3>", unsafe_allow_html=True)
        else:
            st.markdown(f"<h1 style='text-align: center;'>{grupo}</h1>", unsafe_allow_html=True)
            st.markdown(f"<h2 style='text-align: center;'>{ejercicio}</h2>", unsafe_allow_html=True)
            st.markdown(f"<p style='text-align: center;'>{instruccion}</p>", unsafe_allow_html=True)
        
        progress = st.progress(0)
        tiempo_texto = st.empty()
        
        if st.button("Pausar/Reanudar", key="pausar_reanudar"):
            st.session_state.timer_running = not st.session_state.timer_running
    
    if st.session_state.timer_running:
        progress.progress(1 - st.session_state.tiempo_restante / duracion)
        tiempo_texto.markdown(f"<h3 style='text-align: center;'>{st.session_state.tiempo_restante} segundos</h3>", unsafe_allow_html=True)
        
        time.sleep(1)
        st.session_state.tiempo_restante -= 1
        
        if st.session_state.tiempo_restante < 0:
            st.session_state.ejercicio_actual += 1
            if st.session_state.ejercicio_actual < len(rutina):
                st.session_state.tiempo_restante = rutina[st.session_state.ejercicio_actual][2]
            else:
                st.success("¡Felicidades! Has completado tu rutina.")
                if st.button("Volver al inicio", key="volver_inicio"):
                    for key in list(st.session_state.keys()):
                        del st.session_state[key]
                    st.rerun()
                return
        
        st.rerun()

def spotify_link(text, url):
    return st.markdown(f"🎵 [{text}]({url})")
            
def generar_rutina_interface():
    if 'rutina_generada' not in st.session_state:
        st.session_state.rutina_generada = False
    
    if 'mostrar_config' not in st.session_state:
        st.session_state.mostrar_config = False

    if not st.session_state.rutina_generada:
        col1, col2 = st.columns([4, 1])
        
        with col1:
            if st.button("Rutina Rápida", key="rutina_rapida"):
                # Usar configuración predeterminada
                duracion = "30 segundos"
                tiempo_descanso = 10
                vueltas = 3
                generar_rutina_con_config(duracion, tiempo_descanso, vueltas)

        with col2:
            if st.button("⚙️", key="configuracion"):
                st.session_state.mostrar_config = not st.session_state.mostrar_config

        if st.session_state.mostrar_config:
            with st.expander("Configuración", expanded=True):
                duracion = st.selectbox(
                    "Duración de los ejercicios:",
                    ["30 segundos", "40 segundos", "Ambos aleatorios"]
                )
                tiempo_descanso = st.number_input(
                    "Tiempo de descanso entre ejercicios (segundos):",
                    min_value=5, max_value=60, value=10, step=5
                )
                vueltas = st.number_input(
                    "Número de vueltas:",
                    min_value=1, max_value=6, value=3
                )
                if st.button("Generar Rutina Personalizada", key="generar_rutina_personalizada"):
                    generar_rutina_con_config(duracion, tiempo_descanso, vueltas)

    else:
        if 'timer_running' not in st.session_state:
            st.image("entrena3.png", use_column_width=True)
            if st.button("Comenzar Rutina", key="comenzar_rutina"):
                st.session_state.timer_running = True
                st.session_state.ejercicio_actual = 0
                st.session_state.tiempo_restante = st.session_state.rutina[0][2]
                st.rerun()

            st.title("Poné Musica ... ")
            st.image("spotify.png", width=200)
            st.markdown("La música se abrirá en la aplicación, debes volver para hacer clic en COMENZAR RUTINA")



            spotify_playlist_url1 = "https://open.spotify.com/intl-es/track/7BExBy99xIVD7moauE290a?si=5d08e19cd3fd4cd2"
            spotify_link("Lista de reproducción Inglés", spotify_playlist_url1)

            spotify_playlist_url2 = "https://open.spotify.com/intl-es/track/1hWpzhGIPOQ7gKz3ut5eVs?si=f3891cf314774ac0"
            spotify_link("Lista de reproducción Latino", spotify_playlist_url2)

        else:
            temporizador(st.session_state.rutina)

        if st.button("Generar nueva rutina"):
            st.session_state.rutina_generada = False
            st.session_state.timer_running = False
            if 'rutina' in st.session_state:
                del st.session_state.rutina
            st.rerun()

def generar_rutina_con_config(duracion, tiempo_descanso, vueltas):
    with st.spinner("Generando rutina personalizada..."):
        time.sleep(2)  # Espera 2 segundos
        st.session_state.rutina = generar_rutina(duracion, tiempo_descanso, vueltas)
        st.session_state.duracion = duracion
        st.session_state.tiempo_descanso = tiempo_descanso
        st.session_state.vueltas = vueltas
        st.session_state.rutina_generada = True
    st.success("¡Rutina generada con éxito!")
    st.rerun()

def main():
    # Crear un menú lateral para seleccionar el tema

    # Mostrar el banner
    mostrar_logo()

    # Inicializar estado si no existe
    if 'selected_action' not in st.session_state:
        st.session_state.selected_action = None

    # Lógica para manejar el clic del botón y actualización inmediata del estado
    if st.session_state.selected_action is None:
        if st.button("Entrenar por link de YouTube"):
            st.session_state.selected_action = "mostrar_link"
            st.rerun()  # Forzar la recarga inmediata para ejecutar la función seleccionada

        if st.button("Generar rutina de entrenamiento"):
            st.session_state.selected_action = "generar_rutina"
            st.rerun()  # Forzar la recarga inmediata para ejecutar la función seleccionada

    # Ejecutar la función seleccionada
    if st.session_state.selected_action == "mostrar_link":
        clear_page()
        mostrar_link()

    elif st.session_state.selected_action == "generar_rutina":
        clear_page()
        generar_rutina_interface()


    


if __name__ == "__main__":
    main()
